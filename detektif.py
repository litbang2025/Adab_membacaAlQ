# detektif_adab_app.py
# Run: streamlit run detektif_adab_app.py

import streamlit as st
import pandas as pd
import numpy as np
import plotly.express as px
from datetime import date, datetime
from sqlalchemy import create_engine, Column, Integer, String, Date, Table, MetaData, Text, select
from sqlalchemy.engine import Engine
from sqlalchemy.exc import OperationalError
from docx import Document
import io

# -------------------------
# Config & Text
# -------------------------
APP_TITLE = "🔎 Detektif Adab — Aplikasi Analisis Adab Membaca Al-Qur'an"
SUBTITLE = "Amati • Analisis • Perbaiki — Praktis untuk guru SD"

COMPETENCIES = [
    "Peserta didik bertambah keimanannya dengan mengetahui adab membaca Al-Quran",
    "Menunjukkan perilaku beradab ketika membaca Al-Quran",
    "Menjelaskan adab membaca Al-Quran",
    "Menyebutkan dalil-dalil tentang adab membaca Al-Quran",
    "Menyebutkan kisah islam tentang adab membaca Al-Quran",
    "Mempraktikkan adab membaca Al-Quran"
]

ADAB_LIST = [
    ("berniat_ikhlas", "Berniat ikhlas sebelum mulai"),
    ("mulut_bersih", "Mulut tampak bersih"),
    ("keadaan_suci", "Dalam keadaan suci (atau sesuai ketentuan)"),
    ("tempat_bersih", "Tempat membaca bersih"),
    ("menghadap_kiblat_duduk", "Menghadap kiblat & duduk tenang"),
    ("taawudz", "Membaca taawudz sebelum memulai"),
    ("bismillah", "Membaca bismillah di awal (kecuali At-Taubah)"),
    ("khusyu", "Tampak khusyu' (fokus & tadabbur)"),
    ("tartil_yes", "Membaca dengan tartil (pelan & berhenti di akhir ayat)"),
    ("memperindah_suara", "Usaha memperindah suara sesuai kemampuan")
]

TARTIL_ASPEKS = [
    ("artikulasi", "Artikulasi & pelafalan huruf"),
    ("kecepatan", "Kecepatan & pengendalian (berhenti di akhir ayat)"),
    ("ekspresi", "Ekspresi & memperindah suara sesuai kemampuan")
]

MOTIVASI_PERTANYAAN = [
    "Saya bersemangat membaca Al-Qur’an setiap hari.",
    "Saya merasa membaca Al-Qur’an itu menyenangkan.",
    "Saya ingin membaca dengan suara yang indah.",
    "Saya berusaha memahami arti ayat yang saya baca.",
    "Saya merasa senang saat guru memberi pujian."
]

HADITHS = [
    "Hiasilah al-Quran dengan suara kalian. (HR. Ahmad, An-Nasa'i).",
    "Siapa yang tidak memperindah suaranya ketika membaca Al-Qur’an, maka ia bukan dari golongan kami. (HR. Abu Daud, Ahmad).",
    "Bacalah Al-Qur’an, sesungguhnya ia akan datang pada hari kiamat memberikan syafaat bagi pembacanya. (HR. Muslim).",
    "Barangsiapa yang membaca satu huruf dari Kitabullah maka dia akan memperoleh satu kebaikan dan satu kebaikan akan dibalas sepuluh kali. (HR. At-Tirmidzi dan Ad-Darimi)."
]

# -------------------------
# Database (SQLite)
# -------------------------
DB_FILE = "detektif_adab.db"
engine: Engine = create_engine(f"sqlite:///{DB_FILE}", connect_args={"check_same_thread": False})
metadata = MetaData()

# Define observations table (if it exists create_all will ignore)
observations = Table(
    "observations", metadata,
    Column("id", Integer, primary_key=True, autoincrement=True),
    Column("date", Date, nullable=False),
    Column("kelas", String, nullable=True),
    Column("week", String, nullable=True),
    Column("pengamat", String, nullable=True),
    Column("siswa", String, nullable=False),
    Column("surat", String, nullable=True),
    Column("ayat", String, nullable=True),
    Column("berniat_ikhlas", Integer, nullable=True),
    Column("mulut_bersih", Integer, nullable=True),
    Column("keadaan_suci", Integer, nullable=True),
    Column("tempat_bersih", Integer, nullable=True),
    Column("menghadap_kiblat_duduk", Integer, nullable=True),
    Column("taawudz", Integer, nullable=True),
    Column("bismillah", Integer, nullable=True),
    Column("khusyu", Integer, nullable=True),
    Column("tartil_yes", Integer, nullable=True),
    Column("memperindah_suara", Integer, nullable=True),
    Column("artikulasi", Integer, nullable=True),
    Column("kecepatan", Integer, nullable=True),
    Column("ekspresi", Integer, nullable=True),
    Column("total_tartil", Integer, nullable=True),
    Column("mot_pre_sum", Integer, nullable=True),
    Column("mot_post_sum", Integer, nullable=True),
    Column("catatan", Text, nullable=True),
    Column("verified", Integer, nullable=True)
)

try:
    metadata.create_all(engine)
except OperationalError as e:
    # If DB creation fails, show a Streamlit error but allow app to continue (read-only)
    st.error("Gagal membuat/terhubung DB: " + str(e))

# -------------------------
# Util functions
# -------------------------
def compute_total_tartil(rec: dict):
    try:
        return int(rec.get("artikulasi") or 0) + int(rec.get("kecepatan") or 0) + int(rec.get("ekspresi") or 0)
    except Exception:
        return None

def insert_observation(record: dict):
    with engine.begin() as conn:
        conn.execute(observations.insert().values(**record))

def fetch_all():
    # returns DataFrame; safe with SQLAlchemy 2.x via r._mapping
    with engine.begin() as conn:
        res = conn.execute(select(observations))
        rows = [dict(r._mapping) for r in res.fetchall()]
    if not rows:
        return pd.DataFrame(columns=[c.name for c in observations.columns])
    df = pd.DataFrame(rows)
    if "date" in df.columns:
        # ensure datetime for plotting; keep date type as datetime.date for display
        df["date"] = pd.to_datetime(df["date"])
    return df

def calc_percent(col, df):
    if col not in df.columns:
        return None
    s = df[col].dropna()
    if s.empty:
        return None
    vals = s.values
    # binary 0/1
    if set(np.unique(vals)).issubset({0,1}):
        return 100.0 * float(vals.mean())
    else:
        # assume scale 1..3 -> convert to percent of max 3
        max_val = vals.max()
        if max_val <= 3:
            return 100.0 * (float(vals.mean()) / 3.0)
        return 100.0 * float(vals.mean())

def generate_word_report(df: pd.DataFrame, kelas: str, period_desc: str) -> bytes:
    doc = Document()
    doc.add_heading(f"Lap. Detektif Adab — Kelas {kelas}", level=1)
    doc.add_paragraph(f"Periode: {period_desc}")
    doc.add_paragraph("Kompetensi (ringkasan):")
    for c in COMPETENCIES:
        doc.add_paragraph("- " + c, style='List Bullet')
    doc.add_heading("Ringkasan KPI", level=2)
    total_obs = len(df)
    doc.add_paragraph(f"Total observasi: {total_obs}")
    if total_obs:
        for key, label in ADAB_LIST:
            pct = calc_percent(key, df.dropna(subset=[key])) if key == "bismillah" else calc_percent(key, df)
            doc.add_paragraph(f"{label}: {pct:.1f}% " if pct is not None else f"{label}: -")
        if "total_tartil" in df.columns:
            avg_tartil = df["total_tartil"].dropna().mean()
            doc.add_paragraph(f"Rata-rata total tartil: {avg_tartil:.2f}")
    doc.add_heading("Top 3 Adab Perlu Perbaikan", level=2)
    if total_obs:
        pct_map = {}
        for key,label in ADAB_LIST:
            s = df[key].dropna() if key in df.columns else pd.Series()
            if not s.empty:
                pct_map[label] = 100.0 * s.mean()
        sorted_items = sorted(pct_map.items(), key=lambda x: x[1])
        for label, pct in sorted_items[:3]:
            doc.add_paragraph(f"{label} — {pct:.1f}% terpenuhi", style='List Bullet')
    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.read()

# -------------------------
# Streamlit UI — Layout
# -------------------------
st.set_page_config(page_title="Detektif Adab", layout="wide", initial_sidebar_state="expanded")
# Hero Header
st.markdown(
    f"""
    <div style="background: linear-gradient(90deg,#f7fbff 0%,#eef6f9 100%);
                padding: 18px; border-radius: 12px; margin-bottom: 14px;">
        <div style="display:flex; align-items:center; gap:12px;">
            <div style="font-size:40px;">🔎</div>
            <div>
                <h1 style="margin:0; color:#0b3d91;">{APP_TITLE}</h1>
                <div style="color:#243746; font-size:14px;">{SUBTITLE}</div>
            </div>
        </div>
    </div>
    """, unsafe_allow_html=True
)



# -------------------------
# Sidebar
# -------------------------
with st.sidebar:
    # Logo & Judul
    st.markdown(
        """
        <div style="text-align:center;">
            <img src="https://img.icons8.com/fluency/96/reading.png" width="90">
            <h2 style="margin-bottom:0;">📋 Kontrol & Navigasi</h2>
        </div>
        """,
        unsafe_allow_html=True
    )

    st.subheader("⚡ Pengaturan Cepat")
    
    # Ambil data kelas default dari DB
    df_all = fetch_all()
    kelas_default = "4A"
    if not df_all.empty and 'kelas' in df_all.columns:
        kelas_opts = sorted(df_all['kelas'].dropna().unique())
        if kelas_opts:
            kelas_default = kelas_opts[0]

    # Input kelas & minggu
    kelas = st.text_input("🏫 Kelas (default)", value=kelas_default)
    week = st.text_input("📅 Label Minggu / Tim", value=f"Week-{datetime.now().isocalendar()[1]}")

    # Menu Navigasi dengan ikon
    page = st.radio(
        "📌 Menu",
        ["🔎 Observe", "📊 Dashboard", "🧪 Student Lab", "📁 Reports", "⚙️ Admin", "📚 Panduan"],
        index=1
    )

    # Ringkasan Kompetensi
    st.markdown("---")
    with st.expander("📖 Referensi Materi (Ringkasan)"):
        for c in COMPETENCIES:
            st.write(f"• {c}")

    # Kontrol database
    st.markdown("---")
    st.write("🗄️ **Pengaturan Database**")
    if st.checkbox("Tampilkan kontrol database"):
        st.warning("⚠️ Semua data akan dihapus permanen jika Anda melanjutkan!")
        if st.button("🗑️ Reset Database", type="primary"):
            confirm = st.checkbox("✅ Saya mengerti dan ingin melanjutkan", key="confirm_reset")
            if confirm and st.button("🚨 Konfirmasi Reset DB", key="do_reset", type="primary"):
                with engine.begin() as conn:
                    conn.execute(observations.delete())
                st.success("✅ Database telah direset.")

    st.caption("💾 *Simpan & backup file `detektif.db` secara berkala untuk keamanan data.*")


# -------------------------
# Observe Page
# -------------------------
if page == "Observe":
    st.header("🔎 Observasi Adab — Isi Form")
    if 'last_kelas' not in st.session_state:
        st.session_state['last_kelas'] = kelas

    with st.form("obs_form", clear_on_submit=True):
        left, right = st.columns([1,1])
        with left:
            tgl = st.date_input("Tanggal", value=date.today())
            pengamat = st.text_input("Nama Pengamat (siswa/guru)")
            siswa = st.text_input("Nama Siswa yang diamati")
            surat = st.text_input("Surat (singkat)")
            ayat = st.text_input("Ayat (opsional)")
            kelas_in = st.text_input("Kelas (kosong = sidebar kelas)", value=st.session_state.get('last_kelas', kelas))
            week_in = st.text_input("Minggu (kosong = sidebar minggu)", value=week)
        with right:
            st.markdown("### Checklist Adab (Ya=1 / Tidak=0 / kosong=n.r.)")
            adab_vals = {}
            for key,label in ADAB_LIST:
                if key == "bismillah":
                    val = st.selectbox(label, options=["ya","tidak","n.r."], index=0, key=f"bismillah_{label}")
                    v = 1 if val=="ya" else (0 if val=="tidak" else None)
                else:
                    v = st.radio(label, options=[1,0], index=0, horizontal=True, key=f"rb_{key}")
                adab_vals[key] = v

            with st.expander("Rubrik Tartil (klik untuk buka)"):
                tartil_vals = {}
                for key,label in TARTIL_ASPEKS:
                    tartil_vals[key] = st.selectbox(label, options=[1,2,3], index=1, key=f"t_{key}")

            with st.expander("Kuesioner Motivasi (Pre / Post)"):
                st.caption("Isi sesuai perasaan siswa (1 = sangat tidak setuju ... 5 = sangat setuju)")
                mpre = [st.selectbox(f"Pre: {MOTIVASI_PERTANYAAN[i-1]}", options=[1,2,3,4,5], index=2, key=f"pre_{i}") for i in range(1,6)]
                mpost = [st.selectbox(f"Post: {MOTIVASI_PERTANYAAN[i-1]}", options=[1,2,3,4,5], index=2, key=f"post_{i}") for i in range(1,6)]

            catatan = st.text_area("Catatan & Saran (pujian + 1 saran singkat)")
            verified = st.checkbox("Verifikasi (centang jika guru verifikasi)")

        submitted = st.form_submit_button("Simpan Observasi")
        if submitted:
            if not siswa.strip():
                st.error("Nama siswa wajib diisi.")
            else:
                rec = {
                    "date": tgl,
                    "kelas": kelas_in or kelas,
                    "week": week_in or week,
                    "pengamat": pengamat,
                    "siswa": siswa,
                    "surat": surat,
                    "ayat": ayat,
                    "catatan": catatan,
                    "verified": 1 if verified else 0
                }
                for k,_ in ADAB_LIST:
                    rec[k] = adab_vals.get(k)
                for k,_ in TARTIL_ASPEKS:
                    rec[k] = tartil_vals.get(k)
                rec["total_tartil"] = compute_total_tartil(rec)
                rec["mot_pre_sum"] = sum(mpre)
                rec["mot_post_sum"] = sum(mpost)
                insert_observation(rec)
                st.success(f"Observasi untuk {siswa} tersimpan ✅")
                st.session_state['last_kelas'] = rec['kelas']
                with st.expander("Ringkasan observasi baru (klik untuk lihat)"):
                    st.write(pd.DataFrame([rec]))

# -------------------------
# Dashboard Page
# -------------------------
elif page == "Dashboard":
    st.header("📊 Dashboard Analitik — Ringkasan Kelas")
    df = fetch_all()
    if df.empty:
        st.info("Belum ada data observasi. Silakan ke menu Observe untuk memasukkan data.")
    else:
        dff = df.copy()
        try:
            kelas_filter = st.multiselect("Filter Kelas", options=sorted(df["kelas"].dropna().unique()), default=[st.session_state.get('last_kelas', kelas)])
        except Exception:
            kelas_filter = []
        if kelas_filter:
            dff = dff[dff["kelas"].isin(kelas_filter)]

        # date range
        min_date = dff["date"].min().date() if not dff["date"].isna().all() else date.today()
        max_date = dff["date"].max().date() if not dff["date"].isna().all() else date.today()
        date_range = st.date_input("Rentang Tanggal (mulai, akhir)", value=(min_date, max_date))
        if date_range and len(date_range) == 2:
            start, end = date_range
            dff = dff[(pd.to_datetime(dff["date"]) >= pd.to_datetime(start)) & (pd.to_datetime(dff["date"]) <= pd.to_datetime(end))]

        tab1, tab2, tab3, tab4 = st.tabs(["📊 Ringkasan", "🔥 Heatmap", "📈 Tren", "📋 Data"])

        with tab1:
            col1, col2, col3, col4 = st.columns(4)
            col1.metric("Total Observasi", len(dff))
            col2.metric("Rata-rata Total Tartil", f"{dff['total_tartil'].dropna().mean():.2f}" if not dff['total_tartil'].dropna().empty else "-")
            kh = calc_percent('khusyu', dff)
            col3.metric("% Khusyu'", f"{kh:.1f}%" if kh is not None else "-")
            col4.metric("% Tartil Baik (>=8)", f"{100.0 * (dff['total_tartil']>=8).mean():.1f}%" if 'total_tartil' in dff.columns else "-")

            st.markdown("---")
            st.write("### Ringkasan Adab — Persentase terpenuhi")
            adab_pcts = []
            for key,label in ADAB_LIST:
                if key == "bismillah":
                    pct = calc_percent(key, dff.dropna(subset=[key]))
                else:
                    pct = calc_percent(key, dff)
                pct = 0.0 if pct is None else pct
                adab_pcts.append((label, pct))

            left_col, right_col = st.columns(2)
            half = (len(adab_pcts)+1)//2
            for idx,(label,pct) in enumerate(adab_pcts):
                target = left_col if idx < half else right_col
                target.write(f"**{label}** — {pct:.1f}%")
                target.progress(int(np.clip(pct,0,100)))

            st.markdown("---")
            bar_df = pd.DataFrame({"adab": [l for l,_ in adab_pcts], "pct": [p for _,p in adab_pcts]})
            fig = px.bar(bar_df.sort_values("pct"), x="pct", y="adab", orientation="h", text="pct", labels={"pct":"% terpenuhi","adab":"Adab"})
            fig.update_layout(height=420)
            st.plotly_chart(fig, use_container_width=True)

            # Motivasi pre/post plot
            if "mot_pre_sum" in dff.columns and "mot_post_sum" in dff.columns:
                st.markdown("---")
                st.write("### 📈 Peningkatan Motivasi (Pre vs Post)")
                motivasi_df = dff[["siswa","mot_pre_sum","mot_post_sum"]].dropna()
                if not motivasi_df.empty:
                    motivasi_melt = motivasi_df.melt(id_vars="siswa", var_name="Tahap", value_name="Skor")
                    figm = px.bar(motivasi_melt, x="siswa", y="Skor", color="Tahap", barmode="group", title="Perbandingan Skor Motivasi Pre vs Post")
                    st.plotly_chart(figm, use_container_width=True)
                else:
                    st.info("Belum ada data motivasi (Pre/Post).")

        with tab2:
            st.write("### Heatmap: Siswa × Adab")
            heat_cols = [k for k,_ in ADAB_LIST]
            if all(c in dff.columns for c in heat_cols):
                heat_df = dff.groupby("siswa")[heat_cols].mean().fillna(0)
                if not heat_df.empty:
                    fig2 = px.imshow(heat_df, labels=dict(x="Adab", y="Siswa", color="Rata-rata terpenuhi"), x=[l for _, l in ADAB_LIST], y=heat_df.index)
                    st.plotly_chart(fig2, use_container_width=True)
                else:
                    st.info("Heatmap: data tidak mencukupi.")
            else:
                st.info("Heatmap: kolom adab belum lengkap di data.")

        with tab3:
            st.write("### Trend Rata-rata Total Tartil per Minggu")
            if 'week' in dff.columns and 'total_tartil' in dff.columns:
                trend = dff.groupby("week")["total_tartil"].mean().reset_index().sort_values("week")
                if not trend.empty:
                    fig3 = px.line(trend, x="week", y="total_tartil", markers=True)
                    st.plotly_chart(fig3, use_container_width=True)
                else:
                    st.info("Trend: data tidak mencukupi.")
            else:
                st.info("Trend: kolom 'week' atau 'total_tartil' tidak tersedia.")

        with tab4:
            st.write("### Tabel Observasi (terbaru 50)")
            st.dataframe(dff.sort_values("date", ascending=False).head(50))

# -------------------------
# Student Lab Page
# -------------------------
elif page == "Student Lab":
    st.header("🧪 Student Lab — Eksplorasi Sederhana")
    team = st.text_input("Nama Tim", value="Tim Detektif 1")
    question = st.selectbox("Pilih Pertanyaan Penelitian", options=[
        "Apakah rutinitas 10 menit tartil meningkatkan semangat?",
        "Adab mana yang paling sering dilanggar?",
        "Apakah mendengarkan pujian teman membantu tartil?"
    ])
    if st.button("Jalankan Analisis Cepat"):
        df = fetch_all()
        if df.empty:
            st.warning("Data kosong — tambahkan observasi dulu.")
        else:
            st.success("Data dimuat untuk analisis")
            freq = {label: round(calc_percent(key, df) or 0,1) for key,label in ADAB_LIST}
            st.table(pd.DataFrame.from_dict(freq, orient="index", columns=["% terpenuhi"]))
            st.write("Distribusi Total Tartil")
            fig = px.histogram(df, x="total_tartil", nbins=7)
            st.plotly_chart(fig, use_container_width=True)

# -------------------------
# Reports Page
# -------------------------
elif page == "Reports":
    st.header("📁 Generate Reports & Export")
    df = fetch_all()
    if df.empty:
        st.info("Belum ada data untuk laporan.")
    else:
        st.markdown("Pilih kelas & periode")
        kelas_opts = sorted(df["kelas"].dropna().unique()) if "kelas" in df.columns else []
        if kelas_opts:
            kelas_sel = st.selectbox("Kelas", options=kelas_opts, index=0)
            dfk = df[df["kelas"]==kelas_sel]
            period_label = st.text_input("Label periode", value=f"{week}")
            col1, col2 = st.columns(2)
            with col1:
                if st.button("Generate Word Report (.docx)"):
                    bio = generate_word_report(dfk, kelas_sel, period_label)
                    st.download_button(label="Download Laporan .docx", data=bio, file_name=f"Lap_DetektifAdab_{kelas_sel}_{period_label}.docx")
            with col2:
                if st.button("Export CSV Data Kelas"):
                    csv = dfk.to_csv(index=False).encode('utf-8')
                    st.download_button("Download CSV", data=csv, file_name=f"data_detektifadab_{kelas_sel}.csv", mime="text/csv")

            st.markdown("### Temuan Utama (preview)")
            pct_map = {label: (calc_percent(key, dfk) or 0) for key,label in ADAB_LIST}
            sorted_items = sorted(pct_map.items(), key=lambda x: x[1])
            for label,pct in sorted_items[:3]:
                st.write(f"- {label}: {pct:.1f}% terpenuhi")
        else:
            st.info("Belum ada kelas terdaftar di data.")

# -------------------------
# Admin Page
# -------------------------
elif page == "Admin":
    st.header("🔧 Admin — Manajemen Data & Verifikasi")
    df = fetch_all()
    if df.empty:
        st.info("Data kosong.")
    else:
        st.write(f"Total rows: {len(df)}")
        st.dataframe(df.sort_values("date", ascending=False))
        st.markdown("---")
        unverified = df[df["verified"]!=1] if "verified" in df.columns else pd.DataFrame()
        if not unverified.empty:
            st.write("Observasi belum diverifikasi:")
            st.dataframe(unverified)
            sel_id = st.number_input("Masukkan ID untuk tandai verified", min_value=1, step=1)
            if st.button("Tandai Verified"):
                with engine.begin() as conn:
                    conn.execute(observations.update().where(observations.c.id==int(sel_id)).values(verified=1))
                st.success(f"Observasi ID {sel_id} ditandai verified.")
        del_id = st.number_input("Hapus row ID (jika perlu)", min_value=0, step=1)
        if st.button("Hapus Row"):
            if del_id>0:
                with engine.begin() as conn:
                    conn.execute(observations.delete().where(observations.c.id==int(del_id)))
                st.success(f"Row ID {del_id} dihapus.")

# -------------------------
# Panduan Page
# -------------------------
elif page == "Panduan":
    st.header("📚 Panduan Singkat & Etika")
    st.markdown("Referensi adab (ringkasan):")
    for key,label in ADAB_LIST:
        st.write(f"- {label}")
    st.markdown("Keutamaan singkat:")
    for h in HADITHS:
        st.write("- " + h)
    st.markdown("Etika penelitian & privasi:")
    st.write("- Minta persetujuan orang tua sebelum merekam audio atau data sensitif.")
    st.write("- Aplikasi ini menyimpan observasi berbasis teks; tidak merekam suara.")
    st.write("- Data hanya untuk pembelajaran & perbaikan (akses diatur oleh guru).")

# Footer caption in sidebar
st.sidebar.markdown("---")
st.sidebar.caption("Aplikasi mengacu pada materi 'Adab Membaca Al-Qur'an' (hal.28–36).")
